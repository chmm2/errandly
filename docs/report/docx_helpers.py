"""Typography helpers for the Project-I report.

    python docs/report/build_report.py

Writes docs/report/Errandly-Project-I-Report.docx, following the VIT template's
typography exactly:

    chapter heading   Times New Roman 14, bold, UPPER CASE, line spacing 1.5
    section heading   Times New Roman 13, bold, Title Case, line spacing 1.5
    sub-section       Times New Roman 12, bold italic, Title Case, spacing 1.5
    body              Times New Roman 12, line spacing 1.15

Regenerable on purpose: the figures are produced by figures.py and the prose
lives here, so a change to either can be rebuilt into a fresh document rather
than hand-patched into a binary file.
"""

from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

HERE = pathlib.Path(__file__).parent
FIGS = HERE / "figures"
OUT = HERE / "Errandly-Project-I-Report.docx"

TNR = "Times New Roman"


# ─────────────────────────────────────────────────────────── formatting helpers
def _run(p, text, size=12, bold=False, italic=False, caps=False):
    r = p.add_run(text.upper() if caps else text)
    r.font.name = TNR
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = RGBColor(0, 0, 0)
    return r


def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, italic=False,
         space_after=8, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(space_after)
    if indent is not None:
        pf.left_indent = Inches(indent)
    _run(p, text, size=size, italic=italic)
    return p


def chapter(doc, text, page_break=True):
    if page_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(10)
    _run(p, text, size=14, bold=True, caps=True)
    return p


def section(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    _run(p, text, size=13, bold=True)
    return p


def subsection(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    _run(p, text, size=12, bold=True, italic=True)
    return p


def bullet(doc, label, text=""):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.35)
    if text:
        _run(p, f"{label}: ", size=12, bold=True)
        _run(p, text, size=12)
    else:
        _run(p, label, size=12)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.35)
    _run(p, text, size=12)
    return p


def figure(doc, filename, caption, width=6.3):
    path = FIGS / filename
    if not path.exists():
        raise FileNotFoundError(f"missing figure: {path}  (run figures.py first)")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(12)
    _run(c, caption, size=11, bold=True)


def table(doc, headers, rows, widths=None, font=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        _run(p, h, size=font, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            _run(p, val, size=font)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def centered(doc, text, size=12, bold=False, caps=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(after)
    _run(p, text, size=size, bold=bold, caps=caps, italic=italic)
    return p


def pagebreak(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
