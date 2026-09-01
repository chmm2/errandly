"""Build the BCSE497J Project-I report.

    cd docs/report && python build_report.py

Produces Errandly-Project-I-Report.docx. Regenerable on purpose: the figures
come from figures.py and the prose from the sections_* modules, so a change to
either is rebuilt into a fresh document rather than hand-patched into a binary.

Run figures.py first if the figures directory is empty.
"""

from __future__ import annotations

import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).parent))

from docx import Document
from docx.shared import Inches, Pt

from docx_helpers import TNR
from sections_design import chapter_3, chapter_4
from sections_front import abstract, chapter_1, front_matter, table_of_contents
from sections_lit import chapter_2, chapter_5_references

OUT = pathlib.Path(__file__).parent / "Errandly-Project-I-Report.docx"


def main() -> None:
    doc = Document()

    # Template typography: Times New Roman 12 as the document default, so any
    # paragraph added without an explicit run still matches the specification.
    normal = doc.styles["Normal"]
    normal.font.name = TNR
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15

    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.25)
        s.right_margin = Inches(1.0)

    front_matter(doc)
    abstract(doc)
    table_of_contents(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5_references(doc)

    doc.save(OUT)
    print(f"written: {OUT}")
    print(f"  paragraphs: {len(doc.paragraphs)}   tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
