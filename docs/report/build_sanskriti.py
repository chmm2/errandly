"""Full technical document for the trust-aware matching layer.

    cd docs/report && python build_sanskriti.py

Written to stand alone: someone who has read nothing else about the project
should be able to read this and understand what the three prior papers do, what
they leave unsolved, what was built instead, and how it behaves on the actual
running system.

Every number is taken from the deployed constants or from a measurement on the
running stack, so the document can be checked rather than believed.
"""

from __future__ import annotations

import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from docx_helpers import TNR, _run, body, bullet, figure, section, subsection

OUT = pathlib.Path(__file__).parent / "Errandly-Trust-Layer-Sanskriti.docx"


def title(doc, text, sub=None, author=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, size=18, bold=True)
    if sub:
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_after = Pt(6)
        _run(q, sub, size=11.5, italic=True)
    if author:
        r = doc.add_paragraph()
        r.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r.paragraph_format.space_after = Pt(16)
        _run(r, author, size=11, bold=True)


def quote(doc, text, size=10.5):
    return body(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=size,
                italic=True, indent=0.45)


def table(doc, headers, rows, widths, font=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        _run(c.paragraphs[0], h, size=font, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(1)
            _run(p, str(v).replace("**", ""), size=font,
                 bold=str(v).startswith("**"))
    for r in t.rows:
        for i, w in enumerate(widths):
            r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def paper(doc, tag, citation, did, method, leaves):
    subsection(doc, tag)
    quote(doc, citation, size=10)
    body(doc, "What it does. " + did, size=11)
    body(doc, "How. " + method, size=11)
    body(doc, "What it leaves open. " + leaves, size=11)


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = TNR
    normal.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.9)

    title(doc, "Trust-Aware Social Matching",
          "how Errandly decides who is offered an errand, and why a new "
          "friendship is worth less than an old one",
          "Sanskriti Sajlal  ·  23BCE0832")

    # ══════════════════════════════════════════════════ 1. the problem
    section(doc, "1.  The problem this layer solves")
    body(doc,
         "Errandly is a market between strangers. A requester hands money to "
         "someone they may never have met, and a runner fronts cash at a counter "
         "for someone who may never pay them back. Distance alone is a poor way to "
         "choose between candidates: the nearest available student is not "
         "necessarily the one most likely to actually deliver.", size=11)
    body(doc,
         "On a campus, though, trust is not uniform. Students already know some of "
         "their neighbours, and an errand run by an acquaintance carries "
         "materially less risk than one run by a stranger. That existing structure "
         "is free information the platform can use — if it can represent it, "
         "measure it, and resist being gamed by it.", size=11)
    body(doc,
         "This layer is responsible for all three. It builds a friendship graph, "
         "turns a relationship into a number, and hands that number to matching. "
         "Everything downstream — price checking, collusion detection, the escrow "
         "decisions — operates on errands this layer routed.", size=11)

    # ══════════════════════════════════════════════════ 2. prior work
    section(doc, "2.  What the previous work does")
    body(doc,
         "Three papers define the state of the art for this problem. Each is "
         "reviewed for what it establishes and, more usefully, for the assumption "
         "it leaves standing.", size=11)

    paper(doc, "2.1  Jiang et al. — graph-based trust evaluation",
          "W. Jiang, G. Wang, Md. Z. A. Bhuiyan and J. Wu, “Understanding "
          "graph-based trust evaluation in online social networks: methodologies "
          "and challenges,” ACM Computing Surveys, vol. 49, no. 1, 2016.",
          "Surveys and formalises how trust should be computed between two people "
          "who are not directly connected — trust that must travel along a path "
          "through intermediaries rather than existing as a direct rating.",
          "Compares two families of approach: graph simplification, which reduces "
          "a complex network to a tractable set of paths, and graph analogy, which "
          "treats trust like a flow or a resistance across a network. The central "
          "result the survey establishes is that trust must DECAY along a path: a "
          "friend-of-a-friend-of-a-friend cannot count as much as a direct friend, "
          "and the rate of that decay is a design parameter.",
          "Decay is modelled as a function of DISTANCE alone — how many hops "
          "separate two people. The survey treats an edge as an edge. Nothing in "
          "the framework distinguishes a friendship formed three years ago from "
          "one formed this morning; both contribute identically to a path's "
          "strength.")

    paper(doc, "2.2  Chiou and Tu — social-circle trust in ride-hailing",
          "S.-Y. Chiou and T.-Y. Tu, “A trusted mobile ride-hailing evaluation "
          "system with privacy and authentication,” IEEE Access, vol. 8, "
          "pp. 61929–61942, 2020.",
          "Lets a rider judge a driver by the ratings of people close to that "
          "rider in their own social network, rather than by an anonymous global "
          "average — on the argument that a rating from someone you know is worth "
          "more than a rating from a stranger.",
          "A social-network-based scoring scheme with cryptographic protection of "
          "rater identity, implemented and tested on Android. Closeness in the "
          "social graph raises the weight a rater's opinion carries.",
          "The direction is one-way: closeness only ever INCREASES weight. That is "
          "reasonable for the honest case and exactly backwards for the "
          "adversarial one, where a reputation drawn entirely from one's own "
          "circle is the signature of manipulation rather than a mark of quality. "
          "The paper also stops at rating weight; the social graph never reaches "
          "the assignment decision itself.")

    paper(doc, "2.3  Fu and Liu — trust-aware task allocation",
          "F. Donglai and L. Yanhua, “Trust-aware task allocation in collaborative "
          "crowdsourcing model,” The Computer Journal, vol. 64, no. 6, 2021.",
          "Makes trust a first-class objective in task assignment rather than an "
          "afterthought: tasks go to the worker most likely to complete them "
          "properly, not merely the nearest or cheapest.",
          "The T-Aware model scores a worker's trust from their history on the "
          "platform — completion rates, past performance — and optimises that "
          "jointly against task cost during allocation. Assignment reliability "
          "improves over cost-optimal allocation alone.",
          "Trust comes only from the platform's own record of a worker. There is "
          "no peer graph at all: the model has no notion of who knows whom, so a "
          "new user is indistinguishable from any other new user, and the "
          "considerable trust information already present in a closed community "
          "goes unused.")

    subsection(doc, "2.4  The assumption all three share")
    body(doc,
         "Read together, the three papers establish that trust should inform "
         "assignment, that it should decay with distance, and that a rater's "
         "closeness matters. What none of them models is TIME. In every case a "
         "relationship is a static fact: it either exists or it does not, and if "
         "it exists it counts in full from the instant it is created.", size=11)

    # ══════════════════════════════════════════════════ 3. the gap
    section(doc, "3.  The gap, and why it matters here")
    body(doc,
         "On a campus platform, creating a friendship is free and instant. A "
         "student sends a request, another accepts, and an edge exists. If that "
         "edge immediately confers full trust, then the entire social layer can be "
         "acquired in an afternoon.", size=11)
    body(doc, "The concrete attack has two forms.", size=11)
    bullet(doc, "Targeted",
           "identify a student who posts many errands, befriend them, and collect "
           "a 1500-metre head start on everything they post from that moment on.")
    bullet(doc, "Breadth",
           "the reputation layer weights ratings by DISTINCT rater, so repeating a "
           "rating is worthless. The cheapest remaining evasion is therefore to "
           "recruit eighty friends and have each rate once — which scores exactly "
           "as an honest, genuinely popular runner does.")
    body(doc,
         "The second form is the important one, because it is what the rest of the "
         "fraud work pushes an attacker toward. Every one of those eighty ties has "
         "to be created, and created recently. If new edges carry full weight, the "
         "defence has simply moved the attack rather than stopped it.", size=11)
    body(doc,
         "Stated as a gap: the literature decays trust across graph distance, and "
         "leaves it undecayed across time. On a platform where relationships are "
         "cheap to manufacture, that is the dimension an attacker moves in.",
         size=11)

    # ══════════════════════════════════════════════════ 4. what was built
    doc.add_page_break()
    section(doc, "4.  What was built")
    body(doc,
         "Trust between a requester and a candidate runner is one number in "
         "[0, 1], assembled from three factors and handed to matching, where it is "
         "converted into metres of head start.", size=11)
    figure(doc, "fig_trust_flow.png",
           "Fig. 1. How a single trust score is assembled", width=6.3)

    subsection(doc, "4.1  The graph itself")
    body(doc,
         "Accepted friendships are projected from PostgreSQL into Neo4j as FRIEND "
         "edges, each stamped with the date it was accepted. The graph is a "
         "DERIVED read model — nothing is authoritative there, and it can be "
         "rebuilt from PostgreSQL at any time. Every read degrades to a neutral "
         "value if the graph is unavailable, so an outage makes matching less "
         "socially targeted rather than stopping it.", size=11)
    body(doc,
         "Trust is computed over the shortest path between two students, up to "
         "four hops. Beyond four hops a path exists in principle but means nothing "
         "in practice, so those candidates are treated as strangers.", size=11)

    subsection(doc, "4.2  The formula")
    quote(doc, "trust  =  0.45 ^ (hops − 1)   ×   (1 − penalty)   ×   maturity",
          size=12)
    body(doc, "Each factor answers a different question.", size=11)
    table(doc, ["Factor", "Question it answers", "Behaviour"],
          [["0.45 ^ (hops − 1)", "How far away are they?",
            "A direct friend scores 1.0. Each additional hop retains 45%, so a "
            "friend-of-a-friend is 0.45 and by four hops the contribution is "
            "about 9% — enough to prefer a connected stranger over a complete "
            "one, not enough to outrank someone closer."],
           ["1 − penalty", "Is this group sealed, and does money circulate in it?",
            "Discounts trust flowing through a tightly closed neighbourhood. "
            "Structure alone caps the discount at 0.7; money-flow corroboration "
            "raises the cap to 0.95."],
           ["**maturity", "**How new is the newest link on the path?",
            "**The contribution of this work. A brand-new friendship counts for "
            "0.40 and matures to 1.0 over thirty days."]],
          widths=[1.35, 1.85, 3.6], font=9.5)

    subsection(doc, "4.3  Maturity, in detail")
    body(doc,
         "The weight applied is a linear ramp from a floor of 0.40 at zero days to "
         "1.0 at thirty days, and flat thereafter. Three properties of that choice "
         "are deliberate.", size=11)
    bullet(doc, "It is the NEWEST edge on the path that is measured, not the "
           "average",
           "a chain of trust is only as established as its most recently created "
           "link. A two-year friendship reached through an edge made yesterday "
           "tells you about yesterday.")
    bullet(doc, "There is a floor, not a zero",
           "most new friendships are exactly what they appear to be. A new tie is "
           "discounted, never treated as worthless, because the ordinary case "
           "must not be punished to inconvenience the rare one.")
    bullet(doc, "Old edges are NOT decayed",
           "age makes a friendship stronger evidence, not weaker: a tie that has "
           "survived two years is more likely real than one from last week. "
           "Decaying old edges would penalise the entire ordinary population while "
           "costing an attacker nothing, since rebuilding a stale network is free.")
    body(doc,
         "An unknown edge date counts in full. Edges written before the date field "
         "existed are genuinely old, and treating missing data as suspicious would "
         "penalise the platform's earliest users for the platform's own gap.",
         size=11)
    figure(doc, "fig_trust_curves.png",
           "Fig. 2. The maturity ramp, and its effect on trust at each hop distance",
           width=6.5)

    subsection(doc, "4.4  Tiered offering — why sorting alone is not enough")
    body(doc,
         "A subtlety that took a measurement to discover: ranking candidates by "
         "trust does not, on its own, give a friend the errand. Every candidate is "
         "published to within milliseconds of the others, and the first to accept "
         "wins. In that race the nearest stranger usually taps first, whatever the "
         "sort order said.", size=11)
    body(doc,
         "So an errand is WITHHELD from anyone beyond two social hops for the "
         "first 45 seconds. That withholding, not the sort, is what actually gives "
         "someone you know first refusal. The tier then widens on a timer — to "
         "four hops, then to the whole campus — so a student with no connections "
         "is never stranded. If nobody within two hops is nearby at all, the "
         "errand falls straight through to an open offer: an errand nobody sees is "
         "worse than an errand a stranger takes.", size=11)

    subsection(doc, "4.5  What the user sees")
    body(doc,
         "Trust is also surfaced directly in the interface as a connection badge, "
         "in the style of a professional network: 1st for a direct friend, 2nd and "
         "3rd for the degrees beyond, and R for anyone unreachable. Anything past "
         "three hops is shown as R rather than 4th, because the useful distinction "
         "to a runner scanning a feed is “someone I am connected to” versus “a "
         "stranger”, and a fourth-degree path is in practice the latter.", size=11)

    # ══════════════════════════════════════════════════ 5. worked examples
    section(doc, "5.  How it behaves on the actual app")
    body(doc,
         "Every value below is produced by the deployed constants. The rightmost "
         "column converts trust into what it is actually worth during matching, "
         "where the score is distance minus trust × 1500 metres.", size=11)

    subsection(doc, "5.1  Trust for a clean, open neighbourhood")
    table(doc, ["Relationship", "Hops", "Newest edge", "Trust", "Head start"],
          [["Direct friend, established", "1", "over 30 days", "1.0000", "1500 m"],
           ["**Direct friend, made yesterday", "**1", "**1 day", "**0.4200",
            "**630 m"],
           ["Direct friend, made a week ago", "1", "7 days", "0.5400", "810 m"],
           ["Friend of a friend, established", "2", "over 30 days", "0.4500",
            "675 m"],
           ["**Friend of a friend, newest link 1 day", "**2", "**1 day",
            "**0.1890", "**284 m"],
           ["Three hops, established", "3", "over 30 days", "0.2025", "304 m"],
           ["Four hops, established", "4", "over 30 days", "0.0911", "137 m"],
           ["Stranger", "—", "—", "0.0000", "0 m"]],
          widths=[2.5, 0.6, 1.1, 0.9, 1.0], font=9.5)
    body(doc,
         "The two bold rows are the attack being priced. Befriending a heavy "
         "requester this morning buys 630 metres rather than 1500 — enough to "
         "matter between comparable candidates, not enough to beat a genuinely "
         "nearby stranger. The second bold row is worth reading carefully: a "
         "two-hop path whose newest link is a day old scores 0.189, because the "
         "hop decay (0.45) and the maturity weight (0.42) multiply. Freshness and "
         "distance compound rather than trading off.", size=11)

    subsection(doc, "5.2  Trust when the neighbourhood is suspicious")
    table(doc, ["Situation", "Penalty", "Trust", "Head start"],
          [["Two hops through a sealed intermediary\n"
            "(closure 0.90, no money circulating)", "0.544", "0.2050", "308 m"],
           ["Same, but money circulates in that group\n(circulation 0.85)",
            "0.798", "0.0909", "136 m"],
           ["**Direct friend in a sealed group,\n**no money circulating",
            "**0.000", "**1.0000", "**1500 m"],
           ["Direct friend in a group where money\ncirculates (closure 1.0)",
            "0.594", "0.4062", "609 m"]],
          widths=[3.0, 0.9, 0.9, 1.0], font=9.5)
    body(doc,
         "The third row is a deliberate design decision and the one most likely to "
         "be questioned. Structure alone NEVER discounts a direct friend. A "
         "close-knit friend group is the most ordinary thing on a campus and the "
         "strongest honest trust signal the platform has; penalising it on shape "
         "would punish exactly the students the feature exists to serve. Only "
         "money-flow evidence justifies a discount on a direct friendship — the "
         "conjunction of a sealed group AND value circulating inside it, never "
         "either half alone.", size=11)

    subsection(doc, "5.3  A measurement on the running system")
    body(doc,
         "The following was executed against the live stack — real Neo4j, real "
         "Redis geospatial index, real dispatch path — with Chris as requester, "
         "Ujjwal as a direct friend placed 695 metres away, and an unconnected "
         "student placed 31 metres away.", size=11)
    quote(doc,
          "graph:  friend  hops = 1   trust = 0.499\n"
          "graph:  stranger  hops = none   trust = 0.000\n\n"
          "NORMAL round      rank 0   FRIEND     695 m   score  −53.6\n"
          "                  rank 1   stranger    31 m   score   31.0\n"
          "                  →  offered first: FRIEND")
    body(doc,
         "The friend is more than twenty times further away and still ranks first, "
         "because half a friendship is worth roughly 750 metres of walking. That "
         "is the layer working as designed. It is also, precisely, the behaviour "
         "that the collusion work later had to account for — which is discussed in "
         "the companion document.", size=11)

    # ══════════════════════════════════════════════════ 6. decisions
    doc.add_page_break()
    section(doc, "6.  Design decisions worth defending")

    subsection(doc, "6.1  Closure rather than the clustering coefficient")
    body(doc,
         "The measure of how sealed a neighbourhood is could have been the "
         "standard local clustering coefficient. It was not, and the reason "
         "matters. Clustering divides by degree × (degree − 1), so a small clique "
         "cannot reach a high value however closed it is: a four-person ring "
         "scores about 0.5 while a ten-person hostel block scores 0.8. That ranks "
         "the threat exactly backwards, since the small ring is the likelier "
         "collusion unit.", size=11)
    body(doc,
         "Closure — internal edges over internal plus boundary edges — is "
         "size-independent. What marks a suspicious group is that its edges do not "
         "leave it, whether there are four members or ten.", size=11)

    subsection(doc, "6.2  The penalty is not diluted by degree")
    body(doc,
         "An earlier version divided the closure penalty by the intermediary's "
         "degree, on the reasoning that a well-connected person is less likely to "
         "be part of a ring. This had the perverse effect of exempting exactly the "
         "small tight groups the measure exists to catch, and was removed.", size=11)

    subsection(doc, "6.3  The direct-friendship hole")
    body(doc,
         "The closure penalty originally applied only to an intermediary carrying "
         "a multi-hop path. That leaves a hole precisely where collusion lives: in "
         "a three-person ring every member is one hop from every other, there is "
         "no intermediary at all, and the ranker offered them each other's errands "
         "first at full trust. A separate direct-friendship penalty was added, "
         "gated on money-flow evidence for the reason given in §5.2.", size=11)

    subsection(doc, "6.4  Why these constants")
    table(doc, ["Constant", "Value", "Reasoning"],
          [["Hop decay", "0.45",
            "A friend-of-a-friend should be clearly preferable to a stranger and "
            "clearly inferior to a friend. At 0.45 the four-hop contribution is "
            "~9%, which orders candidates without letting distant connections "
            "dominate."],
           ["Maturity window", "30 days",
            "Longer than any plausible farming sprint, short enough that an "
            "ordinary new student reaches full trust within a month of arriving."],
           ["New-friendship floor", "0.40",
            "Low enough that a same-day edge cannot outrank a genuinely near "
            "stranger; high enough that ordinary new friendships still function."],
           ["Closure knee", "0.55",
            "Below this a neighbourhood is not meaningfully sealed. Most honest "
            "friend groups sit below it."],
           ["Structural cap", "0.70",
            "A closed neighbourhood is only CAPABLE of collusion. A full discount "
            "on a suspicion would also make the ranker discontinuous."],
           ["Corroborated cap", "0.95",
            "Never 1.0 — even a corroborated ring may contain someone who simply "
            "has friends, and 1.0 would erase them from matching entirely."],
           ["Social weight", "1500 m",
            "A direct friend beats a stranger up to 1.5 km closer. Tuned so the "
            "graph reorders realistic candidate sets without sending an errand "
            "across campus."]],
          widths=[1.4, 0.75, 4.5], font=9)
    body(doc,
         "All seven are reasoned design constants, not fitted parameters. They "
         "bound the attack rather than model friendship, and fitting them against "
         "real campus data is Project-II work.", size=11)

    # ══════════════════════════════════════════════════ 7. limits
    section(doc, "7.  Limitations, stated plainly")
    bullet(doc, "The constants are unfitted",
           "no live data yet supports the specific values of 30 days, 0.40 or "
           "0.45. They are defensible choices, not measured optima.")
    bullet(doc, "Maturity delays an attack, it does not prevent one",
           "an attacker patient enough to build a network a month in advance pays "
           "no penalty. The mechanism raises the cost and the planning horizon; it "
           "does not close the door.")
    bullet(doc, "The graph can drift",
           "it is derived and rebuildable, but nothing currently detects "
           "divergence from PostgreSQL automatically. A rebuild command exists and "
           "must be run deliberately.")
    bullet(doc, "Trust is symmetric",
           "friendship is modelled as an undirected edge. In reality trust is "
           "often asymmetric — A may rely on B more than B relies on A — and the "
           "current model cannot express that.")

    # ══════════════════════════════════════════════════ 8. crib
    section(doc, "8.  Anticipated questions")
    table(doc, ["Question", "Answer"],
          [["Why not just use distance?",
            "Distance predicts who can get there, not who will deliver properly. "
            "On a closed campus the social graph is free information about the "
            "second question, and prior work [1]–[3] establishes it improves "
            "assignment outcomes."],
           ["Why 30 days and 0.40?",
            "They bound the attack rather than model friendship. 0.40 is low "
            "enough that a same-day edge cannot outrank a genuinely near stranger; "
            "30 days is longer than any plausible farming sprint. Both are stated "
            "as unfitted design constants."],
           ["Why not decay OLD friendships too?",
            "Because age is evidence of authenticity, not staleness. A tie that "
            "has survived two years is more likely real. Decaying it would punish "
            "the ordinary population while costing an attacker nothing, since "
            "rebuilding a stale network is free."],
           ["Why the newest edge, not the average?",
            "A chain of trust is only as established as its weakest, most recent "
            "link. Averaging would let one old friendship launder a brand-new one."],
           ["Isn't preferring friends unfair to new students?",
            "It would be, if it were absolute. The tier widens on a timer and "
            "falls through immediately when nobody connected is nearby, so a "
            "student with no friends is never stranded — they wait 45 seconds."],
           ["What if the graph goes down?",
            "Every read degrades to a neutral value and matching falls back to "
            "distance ordering. Unavailability is never interpreted as evidence "
            "against anyone."],
           ["Is this novel, or just applied?",
            "The graph traversal and hop decay are applied [1]. The novelty is the "
            "time dimension: decaying trust by the age of the newest edge on the "
            "path, which none of the three surveyed approaches models."]],
          widths=[1.9, 4.9], font=9.5)

    section(doc, "References")
    for i, ref in enumerate([
        "W. Jiang, G. Wang, Md. Z. A. Bhuiyan and J. Wu, “Understanding "
        "graph-based trust evaluation in online social networks: methodologies "
        "and challenges,” ACM Computing Surveys, vol. 49, no. 1, pp. 1–35, 2016.",
        "S.-Y. Chiou and T.-Y. Tu, “A trusted mobile ride-hailing evaluation "
        "system with privacy and authentication,” IEEE Access, vol. 8, "
        "pp. 61929–61942, 2020.",
        "F. Donglai and L. Yanhua, “Trust-aware task allocation in collaborative "
        "crowdsourcing model,” The Computer Journal, vol. 64, no. 6, 2021.",
    ], start=1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.15
        pf.space_after = Pt(6)
        pf.left_indent = Inches(0.4)
        pf.first_line_indent = Inches(-0.4)
        _run(p, f"[{i}]\t{ref}", size=10.5)

    doc.save(OUT)
    print(f"written: {OUT}")
    return OUT


if __name__ == "__main__":
    build()
