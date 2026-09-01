"""A plain-language walkthrough of how a collusion ring is caught.

    cd docs/report && python build_walkthrough.py

Produces Errandly-Ring-Detection-Walkthrough.docx (and .pdf if Word is
present). Written to be read on its own during a review, so it repeats context
rather than cross-referencing the main report.

Every number in the worked example is arithmetic anyone can redo by hand: the
point of the document is that the verdict is reproducible, not that it is
clever.
"""

from __future__ import annotations

import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from docx_helpers import TNR, _run, body, bullet, figure, section, subsection


def quote(doc, text, size=10):
    """Verbatim output, left-aligned.

    Justified text stretches a short line of key-value pairs into a row of
    isolated words with rivers of space between them, which reads as broken
    rather than as a quotation.
    """
    return body(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=size,
                italic=True, indent=0.4)

OUT = pathlib.Path(__file__).parent / "Errandly-Ring-Detection-Walkthrough.docx"


def title(doc, text, sub=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, size=18, bold=True)
    if sub:
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_after = Pt(14)
        _run(q, sub, size=11, italic=True)


def step(doc, n, heading):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(4)
    _run(p, f"STAGE {n}   ", size=10, bold=True)
    _run(p, heading, size=13, bold=True)
    return p


def table(doc, headers, rows, widths, font=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        _run(c.paragraphs[0], h, size=font, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(1)
            bold = str(v).startswith("**")
            _run(p, str(v).replace("**", ""), size=font, bold=bold)
    for r in t.rows:
        for i, w in enumerate(widths):
            r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = TNR
    normal.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.9)

    title(doc, "How Errandly Catches a Collusion Ring",
          "a step-by-step walkthrough, with worked numbers")

    body(doc,
         "This document follows one month of activity through the fraud pipeline. "
         "Two groups are traced side by side: an honest friend group and a real "
         "collusion ring. They behave almost identically, and the whole design "
         "exists to separate them.", size=11)

    # ── cast
    section(doc, "The two groups")
    body(doc,
         "Both are triangles of three mutual friends. Both live in the same hostel "
         "block. Both do a lot of errands for one another. On paper they are "
         "indistinguishable — that is the point.", size=11)
    table(doc, ["", "Group A — honest", "Group B — the ring"],
          [["Who", "Three friends who genuinely help each other out.",
            "Three students farming reputation."],
           ["What they do",
            "Take errands when the app offers them, like anyone else.",
            "Post errands for each other that nobody really runs, and pay each "
            "other for them."],
           ["Why",
            "Convenience.",
            "Every completed errand buys a five-star rating and history, at no "
            "real cost. That reputation gets them better ranking later."],
           ["Money",
            "Flows out of the group as often as in.",
            "Goes round the triangle and returns to whoever started it."]],
          widths=[0.9, 2.6, 3.2])

    # ── stage 0
    step(doc, 0, "Every single errand — the offer log")
    body(doc,
         "Nothing about fraud happens here. This stage runs on every errand on the "
         "platform, all month, and simply records what the app did.", size=11)
    body(doc,
         "When an errand is posted, the app finds nearby runners and ranks them. "
         "95% of the time it applies the usual rule, which gives a friend a "
         "1500-metre head start. 5% of the time — chosen at random — it ignores "
         "friendship entirely and ranks on distance alone.", size=11)
    body(doc,
         "Either way it writes one row: who was in the running, what each of them "
         "scored and why, whether friendship was applied on that round, and who "
         "ended up accepting.", size=11)
    body(doc,
         "That 5% is the whole trick. In the other 95% the app decided the answer "
         "in advance, so those errands cannot tell you anything about what people "
         "chose. The 5% are the only errands where nobody had a head start.",
         size=11, italic=True)

    # ── stage 1-2
    step(doc, 1, "The monthly sweep begins — is money going in a circle?")
    body(doc,
         "Once a month a background job looks at the last 180 days of settled "
         "errands and asks whether money forms a closed loop between people who "
         "are also friends: A paid B, B paid C, C paid A.", size=11)
    body(doc,
         "Both of our groups produce such a loop. Friends genuinely do errands for "
         "one another, so this alone accuses nobody.", size=11)

    step(doc, 2, "Is it big enough to bother with?")
    body(doc, "Three floors, all of which both groups clear:", size=11)
    bullet(doc, "At least three people", "two friends paying each other is far "
           "too ordinary to treat as evidence")
    bullet(doc, "At least two full laps", "it happened repeatedly, not once")
    bullet(doc, "Real money per leg", "not a token amount")
    body(doc,
         "At this point the old system would have flagged both groups. The honest "
         "one would have been accused of fraud for doing nothing. That is the "
         "problem the next stage solves.", size=11)

    # ── stage 3
    step(doc, 3, "Did our own router already explain this?")
    body(doc,
         "This is the new stage, and the one that separates the two groups. It "
         "reads the offer log and asks a single question: how much of this "
         "in-group activity did the app itself cause?", size=11)
    body(doc,
         "Take 200 errands posted by members of each group over the window. Ten of "
         "them happened to be exploring rounds.", size=11)

    table(doc,
          ["", "Rounds", "The app expected an in-group take",
           "Group A actually took", "Group B actually took"],
          [["Normal rounds (friend boost on)", "190", "178.6", "177", "**190"],
           ["Exploring rounds (no boost)", "10", "3.0", "3", "**10"],
           ["**Total", "**200", "**181.6", "**180", "**200"]],
          widths=[1.9, 0.7, 1.6, 1.3, 1.3], font=9)

    body(doc,
         "Read the two right-hand columns. In normal rounds both groups behave "
         "roughly as predicted — Group B is a little high, but friends win those "
         "anyway. The difference is in the exploring rounds. With no head start, "
         "Group A takes 3 of 10, about what distance alone predicts. Group B takes "
         "all 10 — grabbing errands where somebody closer was available.", size=11)
    body(doc, "That is a choice, and it is the only thing here that is.", size=11,
         italic=True)

    subsection(doc, "Turning that into a number")
    body(doc,
         "The gap between expected and observed is divided by how much it could "
         "have wobbled by luck. The result is z — a measure of surprise.", size=11)
    table(doc, ["", "Expected", "Observed", "Gap", "z", "Verdict"],
          [["Group A (honest)", "181.6", "180", "−1.6", "−0.4",
            "**explained — no flag"],
           ["Group B (ring)", "181.6", "200", "+18.4", "**5.1",
            "**unexplained — continue"]],
          widths=[1.5, 1.0, 1.0, 0.8, 0.7, 1.8], font=9.5)
    body(doc,
         "Both groups had the same expectation, because both are closed friend "
         "triangles and the router treats them identically. Only what they did "
         "differs.", size=11)
    body(doc,
         "A z below 3 means the gap is small enough to be chance. Group A is "
         "cleared here — the sweep stops, no language model is called, no admin "
         "ever sees them. A z of 3 or more would happen by luck less than once in "
         "a thousand times, so Group B continues.", size=11)
    body(doc,
         "If there are fewer than 25 usable rounds, this stage returns nothing at "
         "all and the sweep behaves exactly as it did before. Too little data must "
         "never produce a confident answer.", size=11)

    # ── stage 4
    step(doc, 4, "The language model — and only now")
    body(doc,
         "Structure and money can tell you who and how much. Neither can tell you "
         "what for. And there is one case where that is the only thing separating "
         "guilt from innocence: three roommates genuinely taking turns fetching "
         "dinner produce exactly the same triangle and exactly the same "
         "circulation as three students farming rewards. The graphs are "
         "identical.", size=11)
    body(doc,
         "What differs is the content. Real errands are messy — different shops, "
         "different items, odd hours, notes written like someone talking to a "
         "friend. Farmed errands are uniform, minimal and oddly regular.", size=11)
    body(doc,
         "So the model is shown Group B's six errands and asked one question: do "
         "these read like real campus life? On the seeded ring it answered:",
         size=11)
    body(doc,
         "reads_as_genuine: false   ·   coherence 0.35   ·   diversity 0.45\n"
         "“Near-identical titles and intervals”   ·   “Single vendor "
         "used for stationery”   ·   “Single vendor used for food”",
         size=10, italic=True, indent=0.4)

    subsection(doc, "Three rules it works under")
    bullet(doc, "Advisory only", "it never punishes, never withholds money and "
           "never raises a severity. It annotates a flag that already exists. Its "
           "useful direction is clearing honest groups, where being wrong is "
           "cheapest.")
    bullet(doc, "The text is written by the accused", "every errand title is "
           "attacker-controlled, so it is fenced, labelled untrusted and the reply "
           "is constrained to a fixed schema. The worst a crafted note can do is "
           "move a number an admin can see and overrule. It can never become an "
           "instruction.")
    bullet(doc, "Silence is neither innocence nor guilt", "if the model is "
           "unavailable or refuses, it returns nothing and everything else behaves "
           "exactly as before.")
    body(doc,
         "It is also asked only about groups the arithmetic already surfaced. "
         "Running it over everybody would cost a fortune and, worse, would put a "
         "language judgement in front of people nothing was wrong with.", size=11)

    # ── stage 5
    step(doc, 5, "One row is saved")
    body(doc,
         "The flag is not a process or a punishment. It is a single row in a "
         "table:", size=11)
    body(doc,
         "person = Vivek   ·   rule = COLLUSION_RING   ·   severity = 3\n"
         "status = OPEN   ·   members = [Vivek, Rohit, Aditya]\n"
         "laps = 2   ·   total = ₹600   ·   z = 5.1   ·   model verdict attached",
         size=10, italic=True, indent=0.4)
    body(doc,
         "One row per member, because a ring has no visible ringleader — the shape "
         "is symmetric, and naming a primary would be inventing a fact.", size=11)

    # ── stage 6
    step(doc, 6, "Three parts of the app read that row")
    body(doc,
         "Nothing else happens automatically. These are not three punishments; "
         "they are three different pieces of code noticing the same row.", size=11)
    table(doc, ["What", "How it works", "Effect"],
          [["They stop being offered each other's errands",
            "Before ranking, matching asks whether the requester shares an open "
            "flag with any candidate, and removes those candidates.",
            "**The circle is broken. This is what actually stops the fraud."],
           ["They drop down the queue for everyone else",
            "Ranking adds up open-flag severities and adds a penalty to the score.",
            "Severity 3 costs 1200 m. They still get work, just less of it."],
           ["It appears in the admin console",
            "The admin page queries for flags whose status is OPEN.",
            "A human is asked to decide."]],
          widths=[1.7, 3.0, 2.1], font=9)
    body(doc,
         "The penalty is capped, and a severity-3 flag costs less than the friend "
         "boost is worth. Suspicion can demote somebody; it can never remove them. "
         "Only a person can make a decision that serious.", size=11)

    # ── stage 7
    step(doc, 7, "A human decides")
    body(doc,
         "The admin sees the money, the laps, the z, and what the model said, and "
         "chooses:", size=11)
    bullet(doc, "Dismiss", "status becomes DISMISSED. Because all three checks "
           "above look only for OPEN or UPHELD flags, they all stop matching at "
           "once — the block lifts, the penalty disappears, it leaves the console. "
           "One word, three effects undone.")
    bullet(doc, "Uphold", "status becomes UPHELD. The effects continue and the "
           "flag keeps counting toward that person's strike level.")

    # ── figure
    doc.add_page_break()
    section(doc, "The whole flow on one page")
    figure(doc, "fig_pipeline.png",
           "Fig. 1. How a collusion flag is raised, end to end", width=6.6)

    # ── crib
    section(doc, "Quick answers")
    table(doc, ["If you are asked…", "Say"],
          [["What is a collusion ring?",
            "Friends posting fake errands for each other and paying each other, so "
            "the money goes round in a circle and comes back. Nobody loses "
            "anything, but each lap manufactures reputation for free."],
           ["Why not just flag friends who transact a lot?",
            "Because our own matcher offers errands to friends first. We would be "
            "flagging people for something we caused. It flags 6 of 6 honest "
            "groups in simulation."],
           ["What is z?",
            "A surprise score. Below 3 the gap could be luck. At 3 or above it "
            "would happen by chance less than once in a thousand times."],
           ["What is the 5% for?",
            "It is the control group. Those are the only errands where friendship "
            "gave nobody an advantage, so they are the only ones that reveal what "
            "people actually chose."],
           ["What does the LLM decide?",
            "Nothing. It answers whether the errands read as genuine, and that "
            "note is attached for the admin. It cannot withhold money or raise a "
            "severity."],
           ["What if the ring sits in the same room?",
            "Harder, and we tested it: still caught at 85% loyalty. They can hide "
            "by restraining themselves to about 60%, but then they are farming at "
            "little more than an honest rate. We rate-limit rings rather than "
            "eliminating them."],
           ["Is this measured or claimed?",
            "The detector and the log are running and tested. The comparison "
            "numbers are from simulation on a synthetic campus. A live pilot is "
            "Project-II."]],
          widths=[2.0, 4.8], font=9.5)

    doc.save(OUT)
    print(f"written: {OUT}")
    return OUT


if __name__ == "__main__":
    build()
