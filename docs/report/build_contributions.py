"""Individual contribution breakdown for the three team members.

    cd docs/report && python build_contributions.py

Deliberately NOT a list of features built. A panel asking "what did each of you
do" is asking which research gap each person owns, what closes it, and why that
is not already in the literature. Implementation detail is evidence for those
claims, not a substitute for them.

The split is one gap per owner, with no shared gap, so that each member can be
questioned separately and answer without deferring to the others.
"""

from __future__ import annotations

import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from docx_helpers import TNR, _run, body, bullet, section, subsection

OUT = pathlib.Path(__file__).parent / "Errandly-Individual-Contributions.docx"

MEMBERS = [
    {
        "name": "Sanskriti Sajlal",
        "reg": "23BCE0832",
        "area": "Trust-Aware Social Matching",
        "papers": "[1] Jiang et al., ACM Computing Surveys 2016  ·  "
                  "[2] Chiou & Tu, IEEE Access 2020  ·  "
                  "[3] Fu & Liu, The Computer Journal 2021",
        "gap":
            "Graph-based trust evaluation [1] models trust decay as a function of "
            "hop distance alone — how far apart two people are in the graph. "
            "Ride-hailing trust scoring [2] weights a rating by the rater's "
            "closeness. Trust-aware task allocation [3] scores trust from platform "
            "history only, with no peer graph at all. In every case an edge is an "
            "edge: a friendship formed years ago and one formed this morning carry "
            "identical weight.",
        "why_it_matters":
            "That is the cheapest attack available on a campus platform. Creating "
            "a friendship costs nothing and takes seconds, so if a fresh edge "
            "confers full trust, the entire social layer can be bought in an "
            "afternoon — befriend a high-volume requester, immediately collect "
            "their errands.",
        "contribution":
            "Trust as a maturity-weighted path length. Trust between two students "
            "is computed over friendship paths up to three hops, and every path is "
            "additionally weighted by the AGE OF ITS NEWEST EDGE: a chain of trust "
            "is only as established as its most recently created link. A brand-new "
            "friendship carries 40% of full weight and matures over thirty days. "
            "Tiered offering then withholds an errand from beyond two hops for the "
            "first 45 seconds, which is what actually gives someone you know first "
            "refusal — merely sorting by trust would leave a race the nearest "
            "stranger usually wins.",
        "novel":
            "Time-weighted trust decay in a task-assignment setting. The surveyed "
            "literature decays trust across graph distance; this decays it across "
            "edge age as well, so a relationship must be established before it can "
            "be spent.",
        "defend":
            "Why 30 days and 40%? They bound the attack rather than model "
            "friendship: 40% is low enough that a same-day edge cannot outrank a "
            "genuinely near stranger, and 30 days is longer than any plausible "
            "farming sprint. Both are stated as unfitted design constants.",
    },
    {
        "name": "Ujjwal Gogoi",
        "reg": "23BDS0335",
        "area": "Price Integrity Without Receipts, and Human Review",
        "papers": "[6] Sarpal et al., arXiv 2023  ·  "
                  "[7] Maranzato & Pereira, IEEE LA-Web 2009  ·  "
                  "[5] Escrow-based P2P payment, 2026",
        "gap":
            "Marketplace price-anomaly detection [6] sets a statistical price bound "
            "from historical and nearby prices, which is exactly the right idea "
            "when no receipt exists — but it assumes structured catalogue items "
            "with ONE price. Marketplace fraud detection [7] extracts behavioural "
            "features and scores sellers, but does so retrospectively: the result "
            "never reaches the mechanism that decides who gets work next.",
        "why_it_matters":
            "A campus sells the same item at materially different prices in "
            "different outlets — the same puff is Rs23 at one canteen and Rs30 at "
            "another. A single campus-wide reference is then wrong in both "
            "directions at once: it reads an honest runner at the dearer shop as "
            "inflating, while leaving a runner inflating at the cheaper one "
            "comfortably inside the band. Meanwhile a runner under active "
            "suspicion keeps receiving exactly as many offers as before.",
        "contribution":
            "A per-store reference price, shrunk toward the campus median in "
            "proportion to how much independent evidence supports it: a shop with "
            "two observations is judged mostly by campus norms, one with twenty "
            "mostly on its own history. Robust estimation throughout — median of "
            "medians with MAD outlier rejection, and one vote per runner so a "
            "single person repeating a claim cannot move the estimate. The "
            "allowance scales with item value, because a flat rupee tolerance "
            "leaves cheap items — most of the traffic — effectively unprotected. "
            "Verdicts feed escrow directly: a flagged claim settles at the "
            "reference and the excess is withheld pending review, so inflating a "
            "price delays money rather than producing it. The administrative "
            "console then lets a human uphold or dismiss, and dismissal restores "
            "the disputed claim as evidence, since a claim judged wrongly should "
            "not go on distorting what the system believes an item costs.",
        "novel":
            "Store-conditioned price consensus with evidence-proportional "
            "shrinkage, and detection wired back into matching as a bounded "
            "ranking penalty rather than a retrospective report.",
        "defend":
            "Why not a separate reference per shop? Splitting outright fragments "
            "small samples until almost everything falls back to no-reference. "
            "Partial pooling keeps every shop usable from its first observation. "
            "And the defence against a shop's reference being walked upward is "
            "independence — distinct runners, not distinct claims.",
    },
    {
        "name": "Chris Martin Mattam",
        "reg": "23BCE0743",
        "area": "Collusion Rings, and Evidence Conditioned on the Router",
        "papers": "[4] Cheng, Chen & Ye, IEEE ICDE 2019  ·  "
                  "[8] Kamvar et al., WWW 2003  ·  "
                  "[9] Viswanath et al., ACM SIGCOMM 2010  ·  "
                  "[10] Perdomo et al., ICML 2020",
        "gap":
            "EigenTrust [8] is the canonical reputation algorithm and its authors "
            "state plainly that it is vulnerable to collusive groups who rate one "
            "another up; the weakness is acknowledged, not solved. Social-graph "
            "defences [9] are shown to be community detection whose effectiveness "
            "rests entirely on the graph being a faithful observation. "
            "Cooperation-aware assignment [4] deliberately co-assigns people who "
            "have worked together before — and never examines that the policy "
            "thereby decides who transacts with whom.",
        "why_it_matters":
            "Put those together and a platform reads its own routing decisions "
            "back as evidence of fraud. Our matcher promotes friends, so friends "
            "transact; our detector reads that concentration as a ring. The "
            "distortion is worst for exactly the tight-knit groups the detector "
            "watches most closely, so the two signals it treats as corroborating "
            "one another — closure and circulation — are correlated through the "
            "routing policy rather than through fraud. In simulation the naive "
            "measure flags six of six entirely honest friend groups.",
        "contribution":
            "Three parts. First, ring detection proper: a ring is identified as a "
            "directed payment cycle among mutual friends, found with Tarjan's "
            "algorithm and gated on group size, laps and per-leg value, with a "
            "local language model asked afterwards — advisory only — whether the "
            "errands read as genuine campus activity, since structure and money "
            "answer who and how much but never what for. Second, the offer log: "
            "every dispatch round records its full candidate set and every ranking "
            "term, because a composite score cannot be decomposed after the fact "
            "and the information is otherwise destroyed the moment the errand ends. "
            "Third, the correction: the routing policy itself becomes the null "
            "hypothesis, so only in-group activity the policy cannot explain is "
            "admitted as evidence, expressed as a test statistic rather than a "
            "hand-tuned threshold.",
        "novel":
            "Conditioning fraud evidence on the assignment policy that produced "
            "it. Performative prediction [10] supplies exactly the right apparatus "
            "and has been developed for prediction and pricing, never for a fraud "
            "graph the platform itself generated. Two further results follow: past "
            "a certain trust-boost strength the policy's expectation saturates and "
            "collusion becomes undetectable in principle, and a small fraction of "
            "deliberately socially-blind offers restores detection power at "
            "negligible routing cost.",
        "defend":
            "Does it catch every ring? No — and that is the honest claim. A "
            "co-located ring can hide by restraining itself to roughly an honest "
            "rate, which means it cannot both stay invisible and farm "
            "aggressively. The method rate-limits collusion, and the limit is "
            "computable from the boost strength and the exploration rate.",
    },
]


def title(doc, text, sub=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, size=18, bold=True)
    if sub:
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_after = Pt(14)
        _run(q, sub, size=11, italic=True)


def table(doc, headers, rows, widths, font=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        _run(c.paragraphs[0], h, size=font, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(1)
            # "**" is an emphasis marker in the row data, not content
            _run(p, str(v).replace("**", ""), size=font,
                 bold=str(v).startswith("**"))
    for r in t.rows:
        for i, w in enumerate(widths):
            r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def field(doc, label, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(7)
    pf.left_indent = Inches(0.9)
    pf.first_line_indent = Inches(-0.9)
    _run(p, f"{label}\t", size=10.5, bold=True)
    _run(p, text, size=10.5)
    return p


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = TNR
    normal.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.75)
        s.left_margin = s.right_margin = Inches(0.85)

    title(doc, "Errandly — Individual Contributions",
          "which research gap each member owns, and what closes it")

    body(doc,
         "The work divides by research gap rather than by software module. Each "
         "member owns one gap from the literature survey, the mechanism that "
         "closes it, and the claim to novelty that follows — so each can be "
         "questioned independently and answer without deferring to the others. "
         "No gap is shared.", size=11)

    section(doc, "At a glance")
    table(doc,
          ["Member", "Area", "Gap owned", "Claim to novelty"],
          [["Sanskriti Sajlal\n23BCE0832", "Trust-aware social matching",
            "Trust decays with graph distance, but an edge's age is ignored — a "
            "friendship made this morning counts the same as one made years ago.",
            "**Time-weighted trust decay: a path is only as established as its "
            "newest link."],
           ["Ujjwal Gogoi\n23BDS0335", "Price integrity and human review",
            "Price-anomaly work assumes one price per item; a campus has the same "
            "item at different prices per shop. Detection also never reaches the "
            "matcher.",
            "**Per-store reference with evidence-proportional shrinkage, fed back "
            "into ranking as a bounded penalty."],
           ["Chris Martin Mattam\n23BCE0743", "Collusion rings and policy-"
            "conditioned evidence",
            "Graph fraud detection assumes the trust graph is observed. Here the "
            "platform's own routing produced it, so the detector reads its own "
            "decisions back as evidence.",
            "**The routing policy as the null hypothesis; the detectability limit "
            "it induces, and exploration as the remedy."]],
          widths=[1.15, 1.35, 2.6, 2.0], font=8.5)

    for m in MEMBERS:
        doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, m["name"], size=15, bold=True)
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(12)
        _run(q, f"{m['reg']}   ·   {m['area']}", size=11, italic=True)

        field(doc, "Papers", m["papers"])
        field(doc, "The gap", m["gap"])
        field(doc, "Why it matters", m["why_it_matters"])
        field(doc, "Contribution", m["contribution"])
        field(doc, "Novelty", m["novel"])
        field(doc, "If challenged", m["defend"])

    doc.add_page_break()
    section(doc, "How the three fit together")
    body(doc,
         "The three contributions are sequential rather than parallel, and the "
         "dependency runs in one direction only.", size=11)
    body(doc,
         "Sanskriti's trust layer decides who is offered an errand. Ujjwal's price "
         "layer decides whether what the runner reported is honest, and withholds "
         "money when it is not. Chris's collusion layer operates on the settlement "
         "history the other two produce — and its central finding is that the "
         "first layer systematically biases the evidence available to the third.",
         size=11)
    body(doc,
         "That dependency is worth stating explicitly at review, because it is the "
         "reason the project has a research contribution at all. Had the trust "
         "layer not been built, the flaw would never have appeared; had the "
         "collusion layer not been built, it would never have been noticed.",
         size=11)

    section(doc, "Shared work")
    body(doc,
         "The platform itself — identity and campus verification, errand lifecycle, "
         "escrow ledger, real-time tracking and chat, the mobile and web clients, "
         "and the deployment — was built jointly and is not claimed by any one "
         "member. It is the substrate the three contributions run on rather than a "
         "contribution in its own right.", size=11)

    doc.save(OUT)
    print(f"written: {OUT}")
    return OUT


if __name__ == "__main__":
    build()
